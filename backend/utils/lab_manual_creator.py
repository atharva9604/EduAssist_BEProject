import os
from datetime import datetime
from pathlib import Path
from typing import Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class LabManualCreator:
    """Creates Word documents (DOCX) from structured lab manual data"""
    
    def __init__(self, storage_dir: str = "storage/lab_manuals"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)
    
    def create_lab_manual_docx(self, manual_data: Dict, filename: str = None) -> str:
        """
        Create a Word document from lab manual data.
        
        Args:
            manual_data: Dictionary containing lab manual structure
            filename: Optional filename (will be generated if not provided)
            
        Returns:
            Path to the created DOCX file
        """
        try:
            # Create new document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_subject = "".join(c for c in manual_data.get("subject", "LabManual") if c.isalnum() or c in (' ', '-', '_')).rstrip()
                filename = f"{safe_subject}_{timestamp}.docx"
            
            # Title
            title = doc.add_heading(manual_data.get("subject", "Lab Manual"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Course Code
            course_code = manual_data.get("course_code")
            if course_code and course_code != "N/A":
                code_para = doc.add_paragraph(f"Course Code: {course_code}")
                code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                code_para.runs[0].font.size = Pt(12)
            
            doc.add_paragraph()  # Spacing
            
            # Prerequisites
            prerequisites = manual_data.get("prerequisites", "")
            if prerequisites:
                doc.add_heading("Prerequisites", level=1)
                prereq_para = doc.add_paragraph(prerequisites)
                prereq_para.runs[0].font.size = Pt(11)
                doc.add_paragraph()  # Spacing
            
            # Lab Objectives
            lab_objectives = manual_data.get("lab_objectives", [])
            if lab_objectives:
                doc.add_heading("Lab Objectives", level=1)
                for obj in lab_objectives:
                    if obj.strip():
                        obj_para = doc.add_paragraph(obj.strip(), style='List Bullet')
                        obj_para.runs[0].font.size = Pt(11)
                doc.add_paragraph()  # Spacing
            
            # Lab Outcomes
            lab_outcomes = manual_data.get("lab_outcomes", [])
            if lab_outcomes:
                doc.add_heading("Lab Outcomes", level=1)
                doc.add_paragraph("At the end of the course, the students will be able to:")
                for outcome in lab_outcomes:
                    if outcome.strip():
                        outcome_para = doc.add_paragraph(outcome.strip(), style='List Bullet')
                        outcome_para.runs[0].font.size = Pt(11)
                doc.add_paragraph()  # Spacing
            
            # Suggested List of Experiments
            doc.add_heading("Suggested List of Experiments", level=1)
            doc.add_paragraph()  # Spacing
            
            # Modules
            modules = manual_data.get("modules", [])
            for module in modules:
                module_num = module.get("module_number", 0)
                module_title = module.get("module_title", f"Module {module_num}")
                selection_req = module.get("selection_requirement", "All")
                
                # Module heading
                if selection_req == "All":
                    module_heading = f"Module {module_num}: {module_title}"
                else:
                    module_heading = f"Module {module_num} ({selection_req}): {module_title}"
                
                doc.add_heading(module_heading, level=2)
                
                # Experiments
                experiments = module.get("experiments", [])
                for exp in experiments:
                    exp_num = exp.get("experiment_number", 0)
                    exp_title = exp.get("title", "")
                    exp_objective = exp.get("objective", "")
                    exp_description = exp.get("description", "")
                    
                    # Experiment number and title
                    exp_heading = doc.add_heading(f"Experiment {exp_num}: {exp_title}", level=3)
                    
                    # Objective
                    if exp_objective:
                        obj_para = doc.add_paragraph(f"Objective: {exp_objective}")
                        obj_para.runs[0].font.size = Pt(11)
                        obj_para.runs[0].italic = True
                    
                    # Description
                    if exp_description:
                        desc_para = doc.add_paragraph(exp_description)
                        desc_para.runs[0].font.size = Pt(11)
                    
                    doc.add_paragraph()  # Spacing between experiments
                
                doc.add_paragraph()  # Spacing between modules
            
            # Save document
            file_path = self.storage_dir / filename
            doc.save(str(file_path))
            
            return str(file_path)
            
        except Exception as e:
            print(f"Error creating lab manual document: {e}")
            raise

